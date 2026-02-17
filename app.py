import streamlit as st
import os
import tempfile
from utils import extract_text_and_images
from ppt_generator import generate_presentation
import google.generativeai as genai
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time
import openai
import json
import json
import graphviz
from scipy import stats
from fpdf import FPDF
from usage_tracker import UsageTracker
import requests
import plotly.express as px
import plotly.graph_objects as go
import numpy as np

# Initialize Tracker
tracker = UsageTracker()

def sanitize_text_for_pdf(text):
    """
    Replaces common non-Latin-1 characters with compatible ASCII equivalents.
    FPDF standard fonts only support Latin-1.
    """
    replacements = {
        '\u2013': '-',   # en-dash
        '\u2014': '--',  # em-dash
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2026': '...', # ellipsis
        '\u2022': '*',   # bullet
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    # Final fallback: encode to latin-1, replacing errors with '?'
    return text.encode('latin-1', 'replace').decode('latin-1')

def call_ai_api(prompt, provider, api_key, model_name=None, json_mode=False):
    """
    Unified wrapper for calling Gemini or OpenAI.
    """
    if provider == "Google Gemini":
        genai.configure(api_key=api_key)
        
        # Determine model to use
        target_model = model_name if model_name else "gemini-1.5-flash"
        
        try:
            model = genai.GenerativeModel(target_model)
            
            # Permissive safety settings for medical/educational content
            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
            ]
            
            config = {"response_mime_type": "application/json"} if json_mode else {}
            response = model.generate_content(prompt, generation_config=config, safety_settings=safety_settings)
            
            # Log Usage (Gemini often returns usage_metadata)
            try:
                if response.usage_metadata:
                    tracker.log_usage("Google Gemini", target_model, response.usage_metadata.prompt_token_count, response.usage_metadata.candidates_token_count)
            except:
                pass # Usage tracking optional for Gemini if it fails

            return response.text
        except Exception as e:
            # Fallback logic for 404s or other model errors
            if "404" in str(e) and target_model != "gemini-1.5-flash":
                try:
                    # Try a safe fallback
                    fallback_model = genai.GenerativeModel("gemini-1.5-flash")
                    config = {"response_mime_type": "application/json"} if json_mode else {}
                    # Use same safety settings for fallback
                    response = fallback_model.generate_content(prompt, generation_config=config, safety_settings=safety_settings)
                    return response.text
                except Exception as fallback_e:
                     raise Exception(f"Gemini API Error (Primary & Fallback): {fallback_e}")
            else:
                # Catch empty response error explicitly if possible or just pass full error
                if "valid Part" in str(e):
                     raise Exception(f"Gemini Refused to Answer (Safety/Empty Response). Try rephrasing. Details: {e}")
                raise Exception(f"Gemini API Error: {e}")

    elif provider == "OpenAI":
        client = openai.OpenAI(api_key=api_key)
        target_model = model_name if model_name else "gpt-4o"
        
        try:
            messages = [{"role": "system", "content": "You are a helpful expert assistant."},
                        {"role": "user", "content": prompt}]
            
            kwargs = {
                "model": target_model, 
                "messages": messages,
            }
            if json_mode:
                kwargs["response_format"] = {"type": "json_object"}
            
            response = client.chat.completions.create(**kwargs)
            
            # Log Usage
            if response.usage:
                tracker.log_usage("OpenAI", target_model, response.usage.prompt_tokens, response.usage.completion_tokens)
                
            return response.choices[0].message.content
        except Exception as e:
            raise Exception(f"OpenAI API Error: {e}")
    return ""

def generate_dalle_image(prompt, api_key):
    """
    Generates an image using DALL-E 3 and returns the local path.
    """
    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1024x1024",
            quality="standard",
            n=1,
        )
        
        image_url = response.data[0].url
        
        # Log Usage (1 image)
        tracker.log_usage("OpenAI", "dall-e-3", 0, 0, image_count=1)
        
        # Download Image
        res = requests.get(image_url)
        if res.status_code == 200:
             with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tf:
                 tf.write(res.content)
                 return tf.name
        return None
    except Exception as e:
        print(f"DALL-E Error: {e}")
        return None

st.set_page_config(page_title="Dental Research Assistant", layout="wide")

# Sidebar for configuration
with st.sidebar:
    st.title("🦷 Dental Research Assistant")
    tool_mode = st.radio("Select Tool", [
        "Presentation Generator", 
        "Results Writer (SPSS-like)",
        "LD Generator (Dissertation)",
        "Postgraduate Study Aid",
        "Thesis Generator",
        "Journal Club Presentation"
    ])
    
    st.header("Configuration")
    
    provider = st.radio("Select AI Provider", ["Google Gemini", "OpenAI"])
    
    api_key_input = st.text_input(f"Enter {provider} API Key", type="password")
    
    api_key = api_key_input
    selected_model = None

    if provider == "Google Gemini":
        st.caption("Get key from [Google AI Studio](https://aistudio.google.com/app/apikey)")
        # Try to list models if key is present, else show defaults
        gemini_models = ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-1.0-pro"]
        if api_key:
            try:
                genai.configure(api_key=api_key)
                fetched_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                if fetched_models:
                    gemini_models = fetched_models
            except:
                pass # Fallback to defaults
        selected_model = st.selectbox("Select Gemini Model", gemini_models, index=0)
        
    else: # OpenAI
        st.caption("Get key from [OpenAI Platform](https://platform.openai.com/api-keys)")
        openai_models = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"]
        selected_model = st.selectbox("Select OpenAI Model", openai_models, index=0)

    with st.expander("ℹ️ Help & Troubleshooting"):
        st.markdown("""
        **Usage Tracking:**
        - Costs are *estimates* based on standard pricing.
        - Check official billing for exact numbers.
        """)
        
    st.divider()
    st.subheader("📊 Usage & Costs")
    st.divider()
    st.subheader("💰 Prepaid Balance")
    
    current_balance = tracker.get_balance()
    st.metric("Remaining Credits", f"${current_balance:.4f}")
    
    with st.expander("Admin: Update Balance"):
        new_balance = st.number_input("Set Current Balance ($)", value=current_balance, min_value=0.0, step=1.0)
        if st.button("Update Balance"):
            tracker.set_balance(new_balance)
            st.success("Balance Updated!")
            st.rerun()

    st.divider()
    st.subheader("📊 Usage Statistics")
    
    daily_stats = tracker.get_daily_stats()
    weekly_stats = tracker.get_weekly_stats()
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Today's Usage", f"${daily_stats['cost']:.3f}")
        st.caption(f"{daily_stats['tokens_in'] + daily_stats['tokens_out']} tokens")
    with col2:
        st.metric("7-Day Usage", f"${weekly_stats['cost']:.2f}")
    
    if provider == "OpenAI":
        st.link_button("OpenAI Billing", "https://platform.openai.com/account/billing/overview")
    else:
        st.link_button("Gemini Billing", "https://aistudio.google.com/app/plan_information")

# ==========================================
# TOOL 1: PRESENTATION GENERATOR
# ==========================================
if tool_mode == "Presentation Generator":
    st.header("Presentation Generator")
    st.markdown("Upload PDFs to generate a postgraduate-level presentation.")
    
    slide_count = st.slider("Target Slide Count", min_value=10, max_value=60, value=50)
    uploaded_files = st.file_uploader("Upload PDF Articles", type=["pdf"], accept_multiple_files=True)

    if uploaded_files and api_key:
        if st.button("Generate Presentation"):
            with st.spinner("Processing PDFs..."):
                try:
                    # Create a temporary directory to store uploaded files and extracted images
                    with tempfile.TemporaryDirectory() as temp_dir:
                        pdf_paths = []
                        for uploaded_file in uploaded_files:
                            file_path = os.path.join(temp_dir, uploaded_file.name)
                            with open(file_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            pdf_paths.append(file_path)

                        # Step 1: Extract Text & Images
                        st.text("Extracting content from PDFs...")
                        extracted_data = extract_text_and_images(pdf_paths, temp_dir)
                        
                        if not extracted_data['text']:
                            st.error("Could not extract text from the PDFs. Please check if they are scanned images without OCR.")
                            st.stop()

                        # Step 2: Generate Content with Gemini
                        st.text("Fetching available AI models...")
                        genai.configure(api_key=api_key)
                        
                        # specific content generation model
                        try:
                            available_models = []
                            for m in genai.list_models():
                                if 'generateContent' in m.supported_generation_methods:
                                    available_models.append(m.name)
                            
                            if not available_models:
                                st.error("No suitable models found for your API Key. Please ensure 'Generative Language API' is enabled in your Google Console.")
                                st.stop()
                        except Exception as e:
                             st.error(f"Error fetching models: {e}. Check your API Key.")
                             st.stop()

                        # Let user select model to avoid 404s
                        selected_model_name = st.selectbox("Select AI Model", available_models, index=0)
                        
                        st.text(f"Generating presentation structure with {selected_model_name}...")
                        
                        # Increase limit to support multiple long PDFs (Gemini 1.5 has large context)
                        full_text = extracted_data['text'][:500000]
                        
                        prompt = f"""
                        You are an expert dental professor creating a comprehensive postgraduate-level presentation.
                        
                        **TASK:** 
                        Synthesize the findings from the following MULTIPLE extracted dental/medical articles into a single, cohesive presentation.
                        
                        **REQUIREMENTS:**
                        1. **VISUALS:** Every slide MUST have a suggested image description.
                        2. **REFERENCES:** Every slide MUST have a specific citation/reference at the bottom (e.g., "Author et al., 2023").
                        3. **DEPTH:** Content must be suitable for residents (Postgraduate level).
                        
                        Create a structured presentation with approximately {slide_count} slides covering:
                        - Introduction & Background
                        - Combined Literature Review 
                        - Methodologies 
                        - Key Clinical Findings / Case Reports
                        - Discussion 
                        - Conclusion
                        - References
                        
                        Output the response as a valid JSON object with the following structure:
                        {{
                            "title": "Comprehensive Presentation Title",
                            "slides": [
                                {{
                                    "title": "Slide Title",
                                    "content": ["Bullet point 1", "Bullet point 2", "Detailed bullet point 3"],
                                    "notes": "Speaker notes for this slide",
                                    "reference": "Smith J et al., Journal of Periodontology 2024",
                                    "suggested_image_type": "clinical_photo" OR "radiograph" OR "graph" OR "diagram" 
                                }}
                            ]
                        }}

                        Extracted Text from File(s):
                        {full_text}
                        """

                        try:
                            # Use Unified Wrapper
                            presentation_content = call_ai_api(prompt, provider, api_key, model_name=selected_model, json_mode=True)
                        except Exception as e:
                            st.error(f"Error generating content: {e}")
                            st.stop()


                        # Step 3: Create PPTX
                        st.text("Building PowerPoint file...")
                        ppt_io = generate_presentation(presentation_content, extracted_data['images'], temp_dir)
                        
                        st.success("Presentation generated successfully!")
                        
                        st.download_button(
                            label="Download Presentation",
                            data=ppt_io,
                            file_name="generated_presentation.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation"
                        )

                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")

    elif not api_key:
        st.warning("Please enter your Google Gemini API Key to proceed.")
    else:
        st.info("Please upload PDF files to begin.")

# ==========================================
# TOOL 2: RESULTS WRITER
# ==========================================
elif tool_mode == "Results Writer (SPSS-like)":
    st.header("Results Writer Assistant")
    st.markdown("Upload your Data (Excel) and a Key Article (PDF) to generate a Results section in the same style.")

    data_file = st.file_uploader("Upload Data Sheet (Excel/CSV/SPSS .sav)", type=["xlsx", "xls", "csv", "sav"])
    ref_pdf = st.file_uploader("Upload Reference Article (PDF)", type=["pdf"])
    protocol_file = st.file_uploader("Upload Study Protocol (Optional - Word/PDF)", type=["docx", "pdf"])
    
    custom_instructions = st.text_area("Specific Instructions for Tables/Analysis (Optional)", 
                                       placeholder="e.g., 'Focus Table 1 on Demographics only', 'Compare Group A vs B specifically for Pain scores'")

    if data_file and ref_pdf and api_key:
        # Load Data Preview (To allow export before generating full report)
        try:
            if data_file.name.endswith('.csv'):
                df = pd.read_csv(data_file)
            elif data_file.name.endswith('.sav'):
                import pyreadstat
                # pd.read_spss requires pyreadstat
                df = pd.read_spss(data_file)
            else:
                df = pd.read_excel(data_file)
                
            # --- OPTIONAL: GRAPH GENERATION ---
            generate_plots = st.checkbox("Generate Graphs? (Uncheck for Text/Tables only)", value=False)

            # --- VARIABLE SELECTOR ---
            all_columns = df.columns.tolist()
            selected_vars = st.multiselect("Select Variables to Analyze (Leave empty for All)", options=all_columns)

            
            # --- FEATURE: EXPORT TO SPSS (.sav) ---
            # Useful for students converting Excel -> SAV for IBM SPSS
            with st.expander("🛠️ Data Tools (Convert to SPSS)"):
                st.write("Convert this dataset to SPSS format (.sav) for use in IBM SPSS software.")
                if st.button("Prepare SPSS Download"):
                    import pyreadstat
                    # Create a temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".sav") as tmp:
                        pyreadstat.write_sav(df, tmp.name)
                        tmp_path = tmp.name
                        
                    with open(tmp_path, "rb") as f:
                        spss_data = f.read()
                        
                    st.download_button(
                        label="Download as SPSS (.sav)",
                        data=spss_data,
                        file_name="converted_data.sav",
                        mime="application/x-spss-sav"
                    )
        except Exception as e:
            st.error(f"Error reading data: {e}")

        # --- FEATURE: SPSS MANUAL GUIDE ---
        with st.expander("🎓 How to Run This in IBM SPSS (Step-by-Step Guide)"):
            st.markdown(f"### Customized Instructions for Your Dataset")
            
            # Detect Structure
            spss_group_col = None
            numeric_cols_all = df.select_dtypes(include=['number']).columns.tolist()
            cat_cols_all = df.select_dtypes(include=['object', 'category']).columns.tolist()
            
            # Simple detection (same as analysis logic)
            for col in cat_cols_all:
                if "group" in col.lower() or "grp" in col.lower() or 2 <= df[col].nunique() <= 6:
                    spss_group_col = col
                    break
            
            if spss_group_col:
                groups = df[spss_group_col].unique()
                groups = [g for g in groups if pd.notna(g)]
                num_groups = len(groups)
                
                st.info(f"Detected Grouping Variable: **'{spss_group_col}'** ({num_groups} Groups)")
                
                # 1. T-TEST (2 Groups)
                if num_groups == 2:
                    st.markdown("""
                    **To compare Numeric Variables (Means) between your 2 groups:**
                    1.  Go to **Analyze** > **Compare Means** > **Independent-Samples T Test**.
                    2.  Move your Numeric Variables (e.g., *""" + ", ".join(numeric_cols_all[:3]) + """...*) into the **Test Variable(s)** box.
                    3.  Move **""" + spss_group_col + """** into the **Grouping Variable** box.
                    4.  Click **Define Groups**.
                        - For Group 1, type: **""" + str(groups[0]) + """**
                        - For Group 2, type: **""" + str(groups[1]) + """**
                    5.  Click **Continue** -> **OK**.
                    """)
                    
                # 2. ANOVA (>2 Groups)
                elif num_groups > 2:
                    st.markdown("""
                    **To compare Numeric Variables (Means) across your """ + str(num_groups) + """ groups:**
                    1.  Go to **Analyze** > **Compare Means** > **One-Way ANOVA**.
                    2.  Move your Numeric Variables into the **Dependent List**.
                    3.  Move **""" + spss_group_col + """** into the **Factor** box.
                    4.  Click **Post Hoc** -> Select **Tukey** (standard) -> **Continue**.
                    5.  Click **OK**.
                    """)
                
                # 3. CHI-SQUARE (Categorical)
                if cat_cols_all:
                    st.markdown("""
                    ---
                    **To compare Categorical Variables (Gender, etc.):**
                    1.  Go to **Analyze** > **Descriptive Statistics** > **Crosstabs**.
                    2.  Move **""" + spss_group_col + """** into **Columns**.
                    3.  Move other Categories (e.g., *""" + ", ".join([c for c in cat_cols_all if c != spss_group_col][:3]) + """*) into **Rows**.
                    4.  Click **Statistics** -> Check **Chi-square**.
                    5.  Click **Continue** -> **OK**.
                    """)
            else:
                st.warning("Could not automatically detect a clear Group column to give specific instructions. Generally, use **Analyze > Descriptive Statistics > Frequencies** for summary stats.")

        if st.button("Generate Results Write-up"):
            with st.spinner("Analyzing data, protocol and generating charts..."):
                try:
                    import matplotlib.pyplot as plt
                    import seaborn as sns
                    
                    # Set publication aesthetics
                    sns.set_theme(style="whitegrid", palette="mako")
                    
                    # 1. Process Data (Already loaded above into df)
                    # Just ensure it's fresh if needed, but df is available from outer scope
                    
                    # Calculate descriptive stats
                    desc_stats = df.describe(include='all').to_string()
                    data_head = df.head(10).to_string()
                    
                    # --- PROCESS PROTOCOL ---
                    protocol_text = "No protocol provided."
                    if protocol_file:
                        if protocol_file.name.endswith('.docx'):
                            doc = Document(protocol_file)
                            protocol_text = "\n".join([para.text for para in doc.paragraphs])
                        elif protocol_file.name.endswith('.pdf'):
                            with tempfile.TemporaryDirectory() as temp_dir:
                                p_path = os.path.join(temp_dir, protocol_file.name)
                                with open(p_path, "wb") as f:
                                    f.write(protocol_file.getbuffer())
                                p_data = extract_text_and_images([p_path], temp_dir)
                                protocol_text = p_data['text']

                    # --- GENERATE CUSTOM STYLED CHARTS ---
                    plot_images = []
                    
                    # FILTER COLUMNS BASED ON SELECTION
                    if selected_vars:
                         # Keep only selected columns for analysis
                         # But be careful to keep the Group column if it exists in df but wasn't selected? 
                         # Actually, user should select Group col if they want grouping.
                         # Or we auto-detect group from entire df, then see if it's in selection.
                         
                         analysis_df = df[selected_vars].copy()
                         numeric_cols = analysis_df.select_dtypes(include=['number']).columns.tolist()
                         cat_cols = analysis_df.select_dtypes(include=['object', 'category']).columns.tolist()
                    else:
                         numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                         cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
                    
                    # 1. Smart Group Detection
                    group_col = None
                    # Prioritize columns with "Group" in name
                    for col in cat_cols:
                        if "group" in col.lower() or "grp" in col.lower():
                            group_col = col
                            break
                    # Fallback to low cardinality
                    if not group_col:
                        for col in cat_cols:
                            if 2 <= df[col].nunique() <= 6:
                                group_col = col
                                break
                    # If still none, maybe first cat col?
                    if not group_col and cat_cols:
                         group_col = cat_cols[0]

                    
                    # --- STATISTICAL ANALYSIS (CALCULATE P-VALUES) ---
                    stats_summary = "No statistical analysis performed (no groups detected).\n"
                    if group_col:
                        stats_summary = f"### Statistical Analysis (Grouping by '{group_col}'):\n"
                        groups = df[group_col].unique()
                        groups = [g for g in groups if pd.notna(g)]
                        
                        for col in numeric_cols:
                            try:
                                # Prepare data for test (drop NaNs)
                                group_data = []
                                for g in groups:
                                    group_data.append(df[df[group_col] == g][col].dropna())
                                
                                # Perform Test
                                p_val = None
                                test_name = ""
                                if len(groups) == 2:
                                    # T-Test
                                    test_name = "Independent T-Test"
                                    stat, p_val = stats.ttest_ind(group_data[0], group_data[1], equal_var=False)
                                elif len(groups) > 2:
                                    # ANOVA
                                    test_name = "One-way ANOVA"
                                    stat, p_val = stats.f_oneway(*group_data)
                                
                                if p_val is not None:
                                    sig_label = "(SIGNIFICANT)" if p_val < 0.05 else "(Not Significant)"
                                    stats_summary += f"- **{col}**: p = {p_val:.4f} {sig_label} [{test_name}]\n"
                                    
                            except Exception as e:
                                stats_summary += f"- {col}: Could not calculate p-value ({str(e)})\n"
                        
                        # --- 2. CATEGORICAL ANALYSIS (Chi-Square) ---
                        # Essential for Table 1 demographics (Gender vs Group, etc.)
                        for col in cat_cols:
                            if col == group_col: continue
                            if df[col].nunique() > 10: continue # Skip text dumps
                            try:
                                # Create Crosstab
                                crosstab = pd.crosstab(df[col], df[group_col])
                                # Perform Chi-Square
                                chi2, p, dof, expected = stats.chi2_contingency(crosstab)
                                
                                sig_label = "(SIGNIFICANT)" if p < 0.05 else "(Not Significant)"
                                stats_summary += f"- **{col}**: p = {p:.4f} {sig_label} [Chi-Square]\n"
                            except Exception as e:
                                pass # Skip if errors (e.g., zeros)
                    
                    
                    def save_plot_to_buffer():
                        buf = io.BytesIO()
                        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
                        buf.seek(0)
                        return buf

                    # --- PLOTTING LOGIC (INTERACTIVE & STATIC) ---
                    # Only run if requested (User Checkbox)
                    if not generate_plots:
                         numeric_cols = []
                         cat_cols = []

                    if group_col:
                        pass
                        # Limit to top 6 distinct numeric vars to avoid flooding
                        # for col in numeric_cols[:6]:
                             
                             # 1. INTERACTIVE PLOTLY (Screen)
                             # try:
                                 # Calculate stats for error bars
                                 # summary_df = df.groupby(group_col)[col].agg(['mean', 'std']).reset_index()
                                 
                                 # fig = px.bar(summary_df, x=group_col, y='mean', 
                                 #              error_y='std',
                                 #              color=group_col,
                                 #              title=f"Mean {col} by {group_col}",
                                 #              labels={'mean': f'Mean {col}', group_col: 'Group'},
                                 #              template="plotly_white")
                                 # fig.update_layout(showlegend=False)
                                 # st.plotly_chart(fig, use_container_width=True)
                             # except Exception as e:
                                 # st.warning(f"Could not render interactive chart for {col}: {e}")

                             # 2. STATIC MATPLOTLIB (Report - High Quality)
                             # plt.figure(figsize=(8, 6))
                             # Using PointPlot for Means + Error Bars (Scientific Style) or BarPlot with error bars
                             # Let's do BarPlot with capsize
                             # sns.barplot(x=group_col, y=col, data=df, errorbar='sd', capsize=.1, palette="mako", hue=group_col, legend=False)
                             
                             # Add individual data points (swarmplot/stripplot) for transparency? 
                             # Maybe too cluttered for basic report. Let's stick to clean bars.
                             # sns.stripplot(x=group_col, y=col, data=df, color='black', alpha=0.3, jitter=True)
                             
                             # plt.title(f"Comparison of {col} by {group_col}", fontsize=14, fontweight='bold', pad=15)
                             # plt.ylabel(f"Mean {col}", fontsize=12)
                             # plt.xlabel(f"{group_col}", fontsize=12)
                             # plt.grid(axis='y', linestyle='--', alpha=0.5)
                             
                             # Add P-value annotation if significant? (Advanced)
                             # For now, keep it simple.
                             
                             # plt.tight_layout()
                             # plot_images.append(save_plot_to_buffer())
                             # plt.close()
                    
                    else:
                        # No Group? Just histograms/distributions
                        for col in numeric_cols[:3]:
                             # Plotly
                             fig = px.histogram(df, x=col, nbins=20, title=f"Distribution of {col}", template="plotly_white")
                             st.plotly_chart(fig, use_container_width=True)
                             
                             # Matplotlib
                             plt.figure(figsize=(8, 5))
                             sns.histplot(df[col], kde=True, color="teal")
                             plt.title(f"Distribution of {col}", fontsize=14, fontweight='bold')
                             plt.tight_layout()
                             plot_images.append(save_plot_to_buffer())
                             plt.close()
                             
                    # --- PLOTTING LOGIC (INTERACTIVE & STATIC) ---
                    
                    # Define keywords for time-series detection
                    time_keywords = ['pre', 'post', '1m', '3m', '6m', 'immed']
                    
                    import re
                    def clean_col_name(n):
                         return re.sub(r'_(pre|post|1m|3m|6m|immed).*', '', n.lower())

                    if group_col:
                        # 1. DEMOGRAPHIC / CATEGORICAL ANALYSIS Only
                        # We specifically want to plot Demographics (Categorical) and NOT numeric means as bars.
                        
                        # Define Professional Teal Palette
                        teal_palette = ["#008080", "#20B2AA", "#48D1CC", "#00CED1", "#5F9EA0", "#B0E0E6"]
                        sns.set_palette(sns.color_palette(teal_palette))
                        
                        for col in cat_cols:
                             if col == group_col: continue # Don't plot group vs group
                             if df[col].nunique() > 10: continue # Skip high cardinality text
                             
                             # CROSS-TABULATION (Grouped Bar)
                             try:
                                 # Calculate Counts & Percentages
                                 ct = pd.crosstab(df[col], df[group_col], normalize='index') * 100
                                 ct_counts = pd.crosstab(df[col], df[group_col])
                                 
                                 # Stacked Bar or Clustered? Clustered is better for comparison
                                 # Prepare data for plotting
                                 plot_data = df.groupby([group_col, col]).size().reset_index(name='Count')
                                 # Calculate percentages within group
                                 total_by_group = df.groupby(group_col).size().reset_index(name='Total')
                                 plot_data = plot_data.merge(total_by_group, on=group_col)
                                 plot_data['Percentage'] = (plot_data['Count'] / plot_data['Total']) * 100
                                 
                                 # A. INTERACTIVE PLOTLY
                                 fig = px.bar(plot_data, x=col, y='Percentage', color=group_col, barmode='group',
                                              text=plot_data['Percentage'].apply(lambda x: '{0:1.2f}%'.format(x)),
                                              color_discrete_sequence=teal_palette,
                                              title=f"Demographic: {col} by {group_col}",
                                              template="plotly_white")
                                 st.plotly_chart(fig, use_container_width=True, key=f"cat_plot_{col}")
                                 
                                 # B. STATIC MATPLOTLIB (TEAL STYLE)
                                 plt.figure(figsize=(8, 6))
                                 ax = sns.barplot(x=col, y='Percentage', hue=group_col, data=plot_data, palette=teal_palette)
                                 
                                 # Add Labels
                                 for container in ax.containers:
                                     ax.bar_label(container, fmt='%.1f%%', padding=3, fontsize=9, fontweight='bold')
                                     
                                 plt.title(f"Comparison of {col} by {group_col}", fontsize=14, fontweight='bold', pad=15)
                                 plt.ylabel("Percentage (%)", fontsize=12)
                                 plt.xlabel(col, fontsize=12)
                                 plt.ylim(0, 100) # Percentage scale
                                 plt.grid(axis='y', linestyle='--', alpha=0.5)
                                 plt.tight_layout()
                                 plot_images.append(save_plot_to_buffer())
                                 plt.close()
                                 
                             except Exception as e:
                                 st.warning(f"Could not plot {col}: {e}")

                        # DISABLE NUMERIC MEAN PLOTS (As requested)
                        # The user specifically asked for "Demographic Data" bars/graphs only.
                        # Numeric data (Means) should be in Tables (handled by 'stats_summary' text).
                        
                        # 2. TIME SERIES (Optional - keeping if it's explicitly strictly pre/post)
                        # ... (Commented out to strictly follow 'only demographic graphs' request)
                        # If user wants ONLY demographics, we skip numeric plots entirely.
                        
                    else:
                         # No Group Column - Pie Charts for Demographics
                         for col in cat_cols:
                             if df[col].nunique() > 10: continue
                             
                             # Pie Chart Data
                             counts = df[col].value_counts()
                             
                             # A. PROFESSIONAL PLOTLY DONUT
                             fig = px.pie(values=counts.values, names=counts.index, 
                                          title=f"Distribution of {col}", 
                                          color_discrete_sequence=teal_palette,
                                          hole=0.4) # Donut style
                             fig.update_traces(textposition='inside', textinfo='percent+label')
                             fig.update_layout(showlegend=False, title_x=0.5) 
                             st.plotly_chart(fig, use_container_width=True, key=f"pie_{col}")
                             
                             # B. PUBLICATION-QUALITY MATPLOTLIB DONUT
                             plt.figure(figsize=(7, 7))
                             # Explode slightly if slices are thin? No, keep it clean.
                             wedges, texts, autotexts = plt.pie(counts.values, 
                                                                autopct='%1.1f%%', 
                                                                startangle=90, 
                                                                colors=teal_palette, 
                                                                wedgeprops=dict(width=0.5, edgecolor='w'), # Donut width
                                                                pctdistance=0.75, # Percentage closer to edge
                                                                textprops={'fontsize': 10, 'weight': 'bold', 'color': 'white'})
                             
                             # Legend instead of labels to avoid clutter
                             plt.legend(wedges, counts.index,
                                        title=col,
                                        loc="center left",
                                        bbox_to_anchor=(1, 0, 0.5, 1))
                             
                             plt.setp(autotexts, size=10, weight="bold")
                             plt.title(f"Distribution of {col}", fontsize=14, fontweight='bold', pad=20)
                             plt.tight_layout()
                             
                             plot_images.append(save_plot_to_buffer())
                             plt.close() 
                    
                        # Numeric Mean Plots are disabled as per user request.

                    # Plot: Correlation Heatmap (Always useful)
                    if len(numeric_cols) > 1:
                        heatmap_cols = numeric_cols[:15] 
                        plt.figure(figsize=(10, 8))
                        corr = df[heatmap_cols].corr()
                        annot_size = 8 if len(heatmap_cols) < 10 else 6
                        sns.heatmap(corr, annot=True, cmap='coolwarm', fmt=".2f", 
                                    square=True, linewidths=.5, annot_kws={"size": annot_size})
                        plt.title("Correlation Matrix", fontsize=14, fontweight='bold')
                        plt.xticks(rotation=45, ha='right')
                        plt.yticks(rotation=0)
                        plt.tight_layout()
                        plot_images.append(save_plot_to_buffer())
                        plt.close()

                    
                    # 2. Process Reference PDF
                    with tempfile.TemporaryDirectory() as temp_dir:
                        file_path = os.path.join(temp_dir, ref_pdf.name)
                        with open(file_path, "wb") as f:
                            f.write(ref_pdf.getbuffer())
                        
                        ref_data = extract_text_and_images([file_path], temp_dir)
                        ref_text = ref_data['text'][:200000]

                    # 3. Generate Write-up
                    prompt = f"""
                    You are a professional medical statistician and researcher.
                    
                    **TASK:**
                    Write the "Results" section for a research paper based on the provided DATA SUMMARY below.
                    
                    **REQUIREMENTS:**
                    1. **USE THE PROTOCOL:** Use the provided "Study Protocol" to correctly identify study groups, variables, and the primary/secondary outcomes. Ensure table groupings match the protocol design.
                    2. **FOUR (4) DETAILED TABLES:** Generate at least 4 markdown tables (Demographics, Descriptive, Comparisons, Correlations). **Format tables strictly in APA style (similar to IBM SPSS Output).**
                    3. **CHART REFERENCES:** Refer to "Figure 1", "Figure 2" etc. (Note: Only Demographics/Categorical graphs are generated).
                    4. **STYLE & TONE:** Mimic the "Reference Article Text". Academic, objective, past tense.
                    
                    **STUDY PROTOCOL:**
                    {protocol_text}
                    
                    **DATA SUMMARY (Descriptive Statistics):**
                    {desc_stats}
                    
                    **RAW DATA SNAPSHOT (First 10 rows):**
                    {data_head}
                    
                    **USER INSTRUCTIONS:**
                    {custom_instructions}
                    
                    **CALCULATED P-VALUES (Use these for the 'Summary'):**
                    {stats_summary}
                    
                    **REFERENCE ARTICLE TEXT (For Style):**
                    {ref_text}
                    
                    **FINAL OUTPUT STRUCTURE:**
                    1. **Results Text**: Organized by logical sections (Demographics, Primary Outcome, Secondary Outcomes).
                    2. **Tables**: Insert Markdown tables where appropriate (APA Format).
                    3. **Summary of Findings**: A distinct section at the very end. 
                       - Explicitly state which group performed better based on the provided P-VALUES.
                       - Use phrases like "statistically significant difference (p<0.05)" or "comparable results (p>0.05)".
                    """

                    try:
                        results_text = call_ai_api(prompt, provider, api_key, model_name=selected_model)
                    except Exception as e:
                        st.error(f"Error generating results: {e}")
                        st.stop()
                    
                    # 4. Create Word Doc
                    doc = Document()
                    doc.add_heading('Results', 0)
                    doc.add_paragraph(results_text)
                    
                    # Add generated plots to the Word Doc
                    if plot_images:
                        doc.add_heading('Figures', level=1)
                        for i, img_buf in enumerate(plot_images):
                            doc.add_paragraph(f"Figure {i+1}")
                            doc.add_picture(img_buf, width=Inches(5.0))
                            img_buf.close()
                    
                    # Save to buffer
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.success(f"Results generated with {len(plot_images)} charts!")
                    st.markdown("### Preview")
                    st.markdown(results_text)
                    
                    # Display charts in Streamlit
                    if plot_images:
                         st.markdown("### Generated Charts")
                         # Re-open buffers for display since docx closed them? NO, docx reads them. 
                         # Actually I should have copied them or not closed them. 
                         # Let's regenerate for display or just trust they are in doc.
                         # Better: Don't close explicitly in loop above, rely on GC or seek(0) again.
                         st.info("Charts have been added to the Word Document.")

                    st.download_button(
                        label="Download Results (.docx)",
                        data=doc_io,
                        file_name="results_section.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"Error: {str(e)}")

    elif not api_key:
        st.warning("Please enter your API Key in the sidebar.")
    else:
        st.info("Please upload both the Data Sheet and the Reference PDF.")

# ==========================================
# TOOL 3: LD GENERATOR (LIBRARY DISSERTATION)
# ==========================================
elif tool_mode == "LD Generator (Dissertation)":
    st.header("Library Dissertation (LD) Generator")
    st.markdown("Expand a synopsis into a full-length dissertation document with references.")

    ld_title = st.text_input("Dissertation Title", placeholder="e.g. Artificial Intelligence in Prosthodontics")
    ld_synopsis = st.text_area("Synopsis / Headings Structure (One per line)", height=200, 
                               placeholder="Introduction\nHistory\nClassification\nApplications\nAdvantages & Disadvantages\nConclusion")
    
    col1, col2 = st.columns(2)
    with col1:
        ld_pages = st.slider("Target Number of Pages", min_value=1, max_value=150, value=5)
    with col2:
        st.info(f"Approx. Word Count: {ld_pages * 400} words")

    if ld_title and ld_synopsis and api_key:
        if st.button("Generate Dissertation"):
            st.info("Starting generation... This may take a moment for longer documents.")
            
            # 1. Parse Synopsis
            headings = [h.strip() for h in ld_synopsis.split('\n') if h.strip()]
            
            full_document_text = ""
            doc = Document()
            # Set Document Style: Times New Roman, Size 12, Double Spacing
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            style.paragraph_format.line_spacing = 2.0
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_heading(ld_title, 0)
            
            progress_bar = st.progress(0)
            
            # Adaptive Pacing
            request_delay = 2 # Start with 2 seconds
            stop_generation = False # Flag to stop on fatal errors

            # Reference Tracking
            ref_counter = 1
            master_references = [] # List of strings: "1. Author..."
            
            # Target References Calculation (40-70 range)
            target_total_refs = 40 if ld_pages < 50 else 70

            for index, heading in enumerate(headings):
                if stop_generation:
                    break

                with st.spinner(f"Writing section: {heading}..."):
                    target_words_per_section = int((ld_pages * 400) / len(headings))
                    
                    # Distribute references evenly
                    current_ref_count = ref_counter - 1
                    remaining_refs_budget = target_total_refs - current_ref_count
                    remaining_sections = len(headings) - index
                    
                    # Calculate how many to ask for this time
                    if remaining_sections > 0:
                        refs_for_this_section = max(0, int(remaining_refs_budget / remaining_sections))
                    else:
                        refs_for_this_section = 0
                        
                    # Cap it reasonably (don't ask for 20 in one go unless needed)
                    refs_for_this_section = min(refs_for_this_section, 10)

                    ref_instruction = ""
                    if refs_for_this_section > 0:
                        ref_instruction = f"""
                        - ADD roughly **{refs_for_this_section}** NEW references.
                        - START numbering from: **{ref_counter}**.
                        - You MUST conclude your response with a section strictly labeled below:
                        ### SECTION_REFERENCES
                        {ref_counter}. First new reference...
                        """
                    else:
                        ref_instruction = f"""
                        - Do NOT add new references. You have reached the citation limit.
                        - CITE EXISTING references from [1] to [{current_ref_count}] where appropriate.
                        - Do NOT include a ### SECTION_REFERENCES block.
                        """

                    prompt = f"""
                    You are an expert dental academician writing a Library Dissertation on "{ld_title}".
                    
                    **CURRENT SECTION:** "{heading}"
                    
                    **INSTRUCTIONS:**
                    1. Write a detailed, expansive academic text for this specific section.
                    2. Target Length: Approximately {target_words_per_section} words.
                    3. **FORMATTING:** The output must start with the Heading **{heading}** in Bold.
                    4. **REFERENCES (CRITICAL):** 
                       - Citation Format in Text: Use simple numbers as citations (e.g. "stated by Smith 5"). Do NOT use brackets, parentheses, or superscripts. Just the number 1, 2, 3 in normal text size.
                       {ref_instruction}
                       
                    5. Content must be highly technical, postgraduate level.
                    
                    Do NOT write the Introduction or Conclusion unless this section IS "Introduction" or "Conclusion". Focus ONLY on "{heading}".
                    """
                    
                    max_retries = 3
                    retry_count = 0
                    success = False
                    
                    while not success and retry_count < max_retries:
                        try:
                            # Adaptive sleep
                            if index > 0 or retry_count > 0:
                                time.sleep(request_delay)
                                
                            # Unified API Call
                            full_response = call_ai_api(prompt, provider, api_key, model_name=selected_model)
                            
                            # Parse Response for References
                            if "### SECTION_REFERENCES" in full_response:
                                parts = full_response.split("### SECTION_REFERENCES")
                                section_text = parts[0].strip()
                                refs_text = parts[1].strip()
                                
                                # Process References
                                new_refs = [r.strip() for r in refs_text.split('\n') if r.strip()]
                                master_references.extend(new_refs)
                                
                                # Update counter based on number of new refs found
                                # Heuristic: if we found 3 lines, we assume 3 refs. 
                                # Ideally we'd parse the numbers, but simple counting is safer for now.
                                ref_counter += len(new_refs)
                            else:
                                section_text = full_response
                                # No refs found or malformed
                            
                            # Add to Word Doc
                            doc.add_heading(heading, level=1)
                            
                            # Clean text: Remove the heading if the AI repeated it at the top
                            clean_text = section_text.replace(f"**{heading}**", "").replace(f"#{heading}", "").strip()
                            
                            # Fix: Split into paragraphs to prevent justification "stretching" on last lines
                            paragraphs = clean_text.splitlines()
                            for p in paragraphs:
                                if p.strip():
                                    doc.add_paragraph(p.strip())
                            
                            full_document_text += f"\n\n{section_text}"
                            success = True
                            
                        except Exception as e:
                            err_msg = str(e).lower()
                            
                            # Check for Quota Exceeded (Out of Credits) - specific to OpenAI
                            if "insufficient_quota" in err_msg:
                                st.error("❌ OpenAI API Quota Exceeded. You have run out of credits.")
                                st.info("Note: A 'ChatGPT Plus' subscription does NOT cover the API. You need to add credits at platform.openai.com/billing.")
                                st.warning("Stopping generation. Please switch to 'Google Gemini' or add OpenAI credits.")
                                success = False 
                                stop_generation = True # Fatal error, stop all
                                break 
                            
                            # Check for Invalid Key
                            elif "invalid_api_key" in err_msg or "authentication failed" in err_msg:
                                st.error("❌ Invalid API Key. Please check your settings.")
                                success = False
                                stop_generation = True
                                break
                            
                            # Check for Rate Limits (Too Fast)
                            elif "429" in err_msg or "rate limit" in err_msg:
                                retry_count += 1
                                
                                # Provider-specific handling
                                if provider == "OpenAI":
                                    wait_time = 10 
                                    request_delay += 2 
                                    msg = f"⚠️ OpenAI Rate Limit. Pausing for {wait_time}s... (Pacing: {request_delay}s)"
                                else:
                                    wait_time = 40
                                    request_delay += 5
                                    msg = f"⏳ Free Tier Limit Hit. Pausing for {wait_time}s to cool down... (You don't need to do anything, just wait!)"

                                st.warning(msg)
                                time.sleep(wait_time)
                            else:
                                st.error(f"Error generating section {heading}: {e}")
                                break 
                    
                    if not success:
                         st.error(f"Failed to generate section '{heading}' after retries.")
                
                progress_bar.progress((index + 1) / len(headings))

            # 3. Compile Master References
            if master_references:
                doc.add_heading("References", level=1)
                import re
                for i, ref in enumerate(master_references):
                    # Clean existing numbering (e.g., "1. Author" -> "Author")
                    clean_ref = re.sub(r'^\d+[\.\)]\s*', '', ref).strip()
                    # Add with strict continuous numbering
                    doc.add_paragraph(f"{i+1}. {clean_ref}")

            # Finalize
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            st.success("Dissertation generated successfully!")
            st.download_button(
                label="Download Dissertation (.docx)",
                data=doc_io,
                file_name=f"{ld_title.replace(' ', '_')}_LD.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif not api_key:
        st.warning("Please enter your API Key.")
    else:
        st.info("Enter details to start.")


# ==========================================
# TOOL 4: POSTGRADUATE STUDY AID
# ==========================================
elif tool_mode == "Postgraduate Study Aid":
    st.header("Postgraduate Study Aid")
    st.markdown("Generate concise essays or short notes from PDFs for exam preparation.")
    
    # 1. Configuration
    departments = ["Oral Pathology", "Oral Medicine", "Periodontics", "Prosthodontics", 
                   "Endodontics", "Orthodontics", "Pedodontics", "Oral Surgery", "Public Health Dentistry"]
    
    col1, col2 = st.columns(2)
    with col1:
        selected_dept = st.selectbox("Department", departments)
    with col2:
        question_topic = st.text_input("Enter Question / Topic (OR describe which question to pick from uploaded papers)", 
                                       placeholder="e.g. 'Describe Ameloblastoma' OR 'Solve Q1 from 2019 Paper'")
        
        # New Controls
        answer_length = st.radio("Answer Type", ["Long Answer (Essay - approx 2 pages)", "Short Answer (Short Note - approx 1 page)"])
        include_mindmap = st.checkbox("Include Mind Map / Flowchart", value=True, help="Visualizes the answer as a flowchart.")
        
        # DALL-E Option (OpenAI Only)
        generate_ai_diagram = False
        if provider == "OpenAI":
             generate_ai_diagram = st.checkbox("Generate AI Illustration (DALL-E 3)", help="Generates a custom dental diagram for the topic.", value=False)
        
    qp_files = st.file_uploader("Upload Question Papers (Optional - Bulk ok)", type=["pdf"], accept_multiple_files=True)
    
    target_qp = None
    if qp_files:
        qp_names = [f.name for f in qp_files]
        target_qp_name = st.selectbox("Select Question Paper to Solve", qp_names)
        # Find the actual file object
        target_qp = next((f for f in qp_files if f.name == target_qp_name), None)

    uploaded_files = st.file_uploader("Upload Textbooks/Articles (Source Material - Optional)", type=["pdf"], accept_multiple_files=True)
    
    if api_key and question_topic:
        if st.button("Generate Answer"):
             with st.spinner("Analyzing text and generating study material..."):
                try:
                    # 1. Process PDFs
                    full_text = "NO TEXTBOOKS PROVIDED. USE YOUR INTERNAL KNOWLEDGE BASE."
                    extracted_data = {'images': []}

                    with tempfile.TemporaryDirectory() as temp_dir:
                        # process QP file (Target ONLY)
                        qp_text = ""
                        if target_qp:
                            fp = os.path.join(temp_dir, "QP_" + target_qp.name)
                            with open(fp, "wb") as f: f.write(target_qp.getbuffer())
                            
                            # Extract just this one
                            qp_data = extract_text_and_images([fp], temp_dir)
                            qp_text = qp_data['text'][:50000] # moderate limit
                        
                        # process Source files
                        if uploaded_files:
                            pdf_paths = []
                            for uploaded_file in uploaded_files:
                                file_path = os.path.join(temp_dir, uploaded_file.name)
                                with open(file_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                pdf_paths.append(file_path)
                            
                            extracted_data = extract_text_and_images(pdf_paths, temp_dir)
                            full_text = extracted_data['text'][:300000] # Limit context
                    
                    if not full_text:
                         st.error("No text found in PDFs.")
                         st.stop()
                    
                    # 2. Generate Content
                    prompt = f"""
                    You are a Professor in **{selected_dept}**.
                    
                    **TASK:**
                    Write a comprehensive answer based on the user's request: "{question_topic}".
                    
                    **FORMAT:** {answer_length}
                    **MODE:** {'INCLUDE MIND MAP' if include_mindmap else 'TEXT ONLY'}
                    
                    **CONTEXT:**
                    1. **Question Paper Content** (File: {target_qp.name if target_qp else 'None'}):
                    {qp_text}
                    
                    2. **Source Material (Textbooks)**:
                    {full_text}
                    
                    **INSTRUCTIONS:**
                    - If the user refers to a specific question (e.g. "Q1 from 2019"), LOCATE it in the "Question Paper Content" first.
                    - Then, ANSWER that question using the "Source Material".
                    - **IMPORTANT**: If "Source Material" says "NO TEXTBOOKS PROVIDED", use your own expert dental knowledge to answer the question comprehensively.
                    
                    **REQUIREMENTS:**
                    1. **Structure**: 
                       - Introduction
                       - Valid Headings & Subheadings (Bullet points where appropriate)
                       - Conclusion
                    
                    {f'''
                    **SPECIAL MIND MAP INSTRUCTION:**
                    You MUST provide the output in TWO STRICT PARTS separated by "---SPLIT---":
                    
                    PART 1: GRAPHVIZ DOT CODE
                    - Create a strictly valid Graphviz DOT code for a mindmap/flowchart of this topic.
                    - Wrap it in ```dot ... ```
                    - Use nice shapes (box, ellipse) and clear labels.
                    
                    ---SPLIT---
                    
                    PART 2: TEXTUAL CONTENT
                    ''' if include_mindmap else 'PART 2: TEXTUAL CONTENT'}

                    **TEXTUAL CONTENT REQUIREMENTS:**
                    1. **Length**: {'Approx 800-1000 words (2 Pages)' if "Long Answer" in answer_length else 'Approx 300-400 words (1 Page)'}.
                    2. **Style**: 
                       - Use **CLEAN BULLET POINTS** for key information. Avoid long, congested paragraphs.
                       - **TABLES**: Include a Comparison Table or Classification Table if relevant.
                    3. **Structure**: 
                       - Introduction
                       - Headings & Subheadings
                       - Conclusion

                    2. **References**:
                       - Cite authors in-text as (Author, Year).
                       - Do NOT include a "References" list at the end.
                    3. **Level**: Postgraduate (MDS) exam level. Concise but high-yield points.
                    
                    **PROVIDED TEXT:**
                    (See Context Above)
                    """
                    
                    content = call_ai_api(prompt, provider, api_key, model_name=selected_model)
                    
                    dot_code = ""
                    text_content = content
                    
                    # 3. Handle Mind Map Output
                    if include_mindmap and "---SPLIT---" in content:
                        parts = content.split("---SPLIT---")
                        
                        # Part 1: Graphviz
                        dot_part = parts[0]
                        if "```dot" in dot_part:
                            dot_code = dot_part.split("```dot")[1].split("```")[0].strip()
                        elif "digraph" in dot_part:
                             dot_code = dot_part.strip()

                        # Part 2: Text Content
                        text_content = parts[1].strip()
                        
                        # Render Logic
                        if dot_code:
                            st.subheader("Interactive Mind Map")
                            try:
                                st.graphviz_chart(dot_code)
                            except Exception as e:
                                st.error(f"Could not render chart: {e}")
                                st.code(dot_code) # Fallback
                        
                    # Clean and Sanitize Text
                    clean_content = text_content.replace('**', '').replace('##', '').replace('###', '')

                    # 4. Generate PDF
                    class PDF(FPDF):
                        def header(self):
                            self.set_font('Arial', 'B', 12)
                            self.cell(0, 10, f'PG Study Aid - {selected_dept}', 0, 1, 'C')
                            self.ln(5)
                            
                        def footer(self):
                            self.set_y(-15)
                            self.set_font('Arial', 'I', 8)
                            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

                    pdf = PDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Arial", size=11)
                    
                    # Add Title
                    pdf.set_font("Arial", 'B', 16)
                    safe_title = sanitize_text_for_pdf(f"Topic: {question_topic}")
                    pdf.multi_cell(0, 10, safe_title)
                    pdf.ln(5)
                    
                    # Add Mind Map Image to PDF
                    if dot_code:
                        try:
                            # Render DOT to PNG file
                            src = graphviz.Source(dot_code)
                            src.format = 'png'
                            # Use temp_dir from the outer scope
                            mindmap_path = src.render(filename='mindmap_temp', directory=temp_dir, cleanup=True)
                            
                            pdf.set_font("Arial", 'B', 12)
                            pdf.cell(0, 10, "Mind Map / Flowchart", 0, 1, 'L')
                            # Constrain width
                            pdf.image(mindmap_path, x=10, w=190) 
                            pdf.ln(10)
                            
                        except Exception as e:
                            # Fallback: Outline is not explicitly separated now, user gets text. 
                            print(f"Graphviz binary likely missing: {e}")
                            pdf.set_font("Courier", size=10)
                            pdf.multi_cell(0, 5, "[Mind Map could not be rendered graphically in PDF (Software missing)]")
                            pdf.ln(5)

                    # Add DALL-E Image to PDF
                    if generate_ai_diagram and provider == "OpenAI":
                        try:
                            with st.spinner("Generating AI Illustration (DALL-E 3)..."):
                                dalle_prompt = f"A clear, professional dental illustration of {question_topic}, educational style, white background, detailed anatomy label style."
                                dalle_img_path = generate_dalle_image(dalle_prompt, api_key)
                                
                                if dalle_img_path:
                                    pdf.add_page()
                                    pdf.set_font("Arial", 'B', 14)
                                    pdf.cell(0, 10, "AI Generated Illustration", 0, 1, 'L')
                                    pdf.ln(5)
                                    # Constrain width
                                    pdf.image(dalle_img_path, x=20, w=170)
                                    pdf.set_font("Arial", 'I', 9)
                                    pdf.cell(0, 5, f"Figure: AI Illustration of {question_topic} (DALL-E 3)", 0, 1, 'C')
                                    pdf.ln(10)
                        except Exception as e:
                            st.warning(f"Could not generate DALL-E image: {e}")

                    # Add Content (Standard or Outline)
                    pdf.set_font("Arial", size=11)
                    
                    # Sanitize content
                    safe_content = sanitize_text_for_pdf(clean_content)
                    
                    # Line-by-line processing for Table detection
                    for line in safe_content.split('\\n'):
                        stripped = line.strip()
                        # Simple Table Detection: Starts and ends with |
                        if stripped.startswith('|') and stripped.endswith('|'):
                            pdf.set_font("Courier", size=9) # Monospace for tables
                            pdf.multi_cell(0, 5, line)
                        else:
                            pdf.set_font("Arial", size=11)
                            pdf.multi_cell(0, 6, line)
                    
                    # 5. Append Extracted Images
                    if extracted_data['images']:
                        pdf.add_page()
                        pdf.set_font("Arial", 'B', 14)
                        pdf.cell(0, 10, "Relevant Figures from Text", 0, 1, 'L')
                        pdf.ln(5)
                        
                        # Take up to 5 images
                        for idx, img in enumerate(extracted_data['images'][:5]):
                            try:
                                # Start new page for every 2 images to avoid weird breaks
                                if idx > 0 and idx % 2 == 0:
                                    pdf.add_page()
                                
                                img_path = img['path']
                                # Insert Image (centered, width=150mm)
                                pdf.image(img_path, x=30, w=150)
                                pdf.ln(5)
                                pdf.set_font("Arial", 'I', 9)
                                pdf.cell(0, 5, f"Figure {idx+1} (Source: {img['source']}, Page {img['page']+1})", 0, 1, 'C')
                                pdf.ln(10)
                            except Exception as e:
                                print(f"Error adding image {img}: {e}")
                    
                    # Output
                    pdf_output = pdf.output(dest='S').encode('latin-1')
                    
                    st.success("Study Aid Generated!")
                    
                    st.download_button(
                        label="Download PDF",
                        data=pdf_output,
                        file_name=f"PG_Study_Aid_{selected_dept}.pdf",
                        mime="application/pdf"
                    )

                except Exception as e:
                    st.error(f"Error: {str(e)}")
                    
    elif not api_key:
        st.warning("Please enter your API Key.")
    else:
        st.info("Upload PDFs to begin.")

# ==========================================
# TOOL 5: THESIS GENERATOR
# ==========================================
elif tool_mode == "Thesis Generator":
    st.header("Thesis / Dissertation Generator")
    st.markdown("Generate full thesis chapters based on your protocol and university guidelines.")
    
    # 1. Configuration & Uploads
    st.markdown("### Step 1: Core Context")
    col1, col2 = st.columns(2)
    with col1:
        univ_format = st.file_uploader("University Format/Guidelines (Optional - Multiple PDF/Docx)", type=["pdf", "docx"], accept_multiple_files=True, help="Formatting rules and structure.")
        protocol_files = st.file_uploader("Study Protocol (PDF/Docx/Txt - Multiple)", type=["pdf", "docx", "txt"], accept_multiple_files=True, help="Contains Intro, Aims, Methods.")
    with col2:
        key_articles = st.file_uploader("Key Articles for Review (PDFs)", type=["pdf"], accept_multiple_files=True, help="Source material for Literature Review.")
        ref_chapters = st.file_uploader("Reference Chapter (Optional - Multiple PDF/Docx)", type=["pdf", "docx"], accept_multiple_files=True, help="Example of required reference style.")
    
    st.markdown("### Step 2: Results & Discussion")
    results_files = st.file_uploader("Results Data/Files (Optional - Multiple)", type=["pdf", "docx", "xlsx", "csv"], accept_multiple_files=True, help="Upload results to generate the Discussion chapter.")
    
    generate_discussion = False
    if results_files:
        generate_discussion = True
        st.success("Results uploaded. Discussion chapter will be generated.")
    else:
        st.info("No Results uploaded. Discussion chapter will be skipped.")

    if st.button("Generate Thesis Chapters") and api_key:
        if not protocol_files:
            st.error("Study Protocol is required to generate Introduction and Methods.")
            st.stop()
            
        with st.spinner("Analyzing files and generating chapters..."):
            try:
                # 1. Text Extraction
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Helper to extract text from file obj
                    def get_text(file_obj):
                         if not file_obj: return ""
                         fp = os.path.join(temp_dir, file_obj.name)
                         with open(fp, "wb") as f: f.write(file_obj.getbuffer())
                         
                         if file_obj.name.endswith('.pdf'):
                             ret = extract_text_and_images([fp], temp_dir)
                             return ret['text']
                         elif file_obj.name.endswith('.docx'):
                             try:
                                 d = Document(fp)
                                 return "\n".join([p.text for p in d.paragraphs])
                             except: return ""
                         else: # txt/csv
                             try: return file_obj.getvalue().decode("utf-8")
                             except: return ""

                    # Helper to process LIST of files
                    def get_text_from_list(file_list):
                        full_text = ""
                        if file_list:
                            for f in file_list:
                                full_text += get_text(f) + "\n\n"
                        return full_text

                    protocol_text = get_text_from_list(protocol_files)
                    univ_text = get_text_from_list(univ_format)
                    ref_style_text = get_text_from_list(ref_chapters)
                    results_text = get_text_from_list(results_files)
                    
                    article_texts = get_text_from_list(key_articles)
                    
                    # Truncate for context window
                    protocol_text = protocol_text[:30000]
                    univ_text = univ_text[:10000]
                    article_texts = article_texts[:100000] # Large context
                    results_text = results_text[:30000]
                    
                    # 2. Define Chapters to Generate
                    chapters = ["Introduction", "Aims and Objectives", "Review of Literature", "Materials and Methods"]
                    if generate_discussion:
                        chapters.append("Discussion")
                    
                    # 3. Initialize Doc
                    doc = Document()
                    # Style: TNR 12, Double Spaced, Justified
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    style.paragraph_format.line_spacing = 2.0
                    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    progress_bar = st.progress(0)
                    
                    for i, chapter in enumerate(chapters):
                        with st.spinner(f"Writing {chapter}..."):
                            
                            # Construct Prompt based on Chapter
                            context_instruction = ""
                            if chapter == "Introduction":
                                context_instruction = f"Use the 'Study Protocol' to write a strong introduction. Background: {protocol_text[:5000]}..."
                            elif chapter == "Aims and Objectives":
                                context_instruction = f"Extract and refine Aims from 'Study Protocol': {protocol_text[:5000]}..."
                            elif chapter == "Review of Literature":
                                context_instruction = f"Synthesize the 'Key Articles' into a coherent review. Articles: {article_texts}..."
                            elif chapter == "Materials and Methods":
                                context_instruction = f"Expand the 'Study Protocol' methods section. Protocol Metadata: {protocol_text}..."
                            elif chapter == "Discussion":
                                context_instruction = f"Compare 'Results' ({results_text}) with 'Key Articles' ({article_texts[:10000]}). Explain findings."
                                
                            ref_prompt = "Reference Style: Standard (Author, Year) or Numbered. If 'Reference Chapter' text is provided to you below, MIMIC IT."
                            if ref_style_text:
                                ref_prompt += f"\nREFERENCE SAMPLE: {ref_style_text[:2000]}"
                                
                            prompt = f"""
                            You are an expert academic writer for dental/medical theses.
                            
                            **TASK:** Write the **{chapter}** chapter for a thesis.
                            
                            **UNIVERSITY GUIDELINES (Structure/Format):**
                            {univ_text}
                            
                            **CONTENT SOURCE:**
                            {context_instruction}
                            
                            **INSTRUCTIONS:**
                            1. Write expansive, high-quality academic text.
                            2. **Formatting**: Starts with Heading "{chapter}".
                            3. **References**: {ref_prompt}
                            4. **Tone**: Formal, objective, past tense for methods/results, present/past for review.
                            
                            Write ONLY the content for {chapter}.
                            """
                            
                            response = call_ai_api(prompt, provider, api_key, model_name=selected_model)
                            
                            # Add to Doc
                            doc.add_heading(chapter, level=0)
                            # Remove markdown heading if present
                            clean_text = response.replace(f"# {chapter}", "").replace(f"**{chapter}**", "").strip()
                            doc.add_paragraph(clean_text)
                            doc.add_page_break()
                            
                            progress_bar.progress((i + 1) / len(chapters))
                            
                    # 4. Finalize
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.success("Thesis Chapters Generated Successfully!")
                    st.download_button(
                        label="Download Thesis (.docx)",
                        data=doc_io,
                        file_name="Thesis_Draft.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"Error during generation: {e}")
                
    elif not api_key:
        st.warning("Please enter your API Key.")

# ==========================================
# TOOL 6: JOURNAL CLUB PRESENTATION
# ==========================================
elif tool_mode == "Journal Club Presentation":
    st.header("Journal Club Presentation Generator")
    st.markdown("Convert a research article into a critical analysis presentation (Journal Club format).")
    
    # Files
    jc_article = st.file_uploader("Upload Article (PDF)", type=["pdf"], help="The research paper to present.")
    slide_count = st.slider("Target Slide Count", min_value=10, max_value=30, value=15)
    
    if st.button("Generate Presentation") and api_key:
        if not jc_article:
            st.error("Please upload an article.")
            st.stop()
            
        with st.spinner("Analyzing article and crafting presentation..."):
            try:
                # 1. Extract Text & Images
                with tempfile.TemporaryDirectory() as temp_dir:
                    fp = os.path.join(temp_dir, jc_article.name)
                    with open(fp, "wb") as f: f.write(jc_article.getbuffer())
                    
                    extracted_data = extract_text_and_images([fp], temp_dir)
                    full_text = extracted_data['text'][:80000] # Limit context
                    images_found = extracted_data['images']
                    
                    # Compile Image Metadata for Context
                    image_context = ""
                    if images_found:
                        image_context = "AVAILABLE IMAGES (Use 'suggested_image_index' to insert):\n"
                        for idx, img in enumerate(images_found):
                            image_context += f"Index {idx}: Image on Page {img['page']+1}\n"
                    
                    # 2. Generate JSON Content
                    prompt = f"""
                    You are an expert academic presenter preparing a 'Journal Club' presentation for a Dental/Medical conference.
                    
                    **TASK:** Create a critical analysis presentation for the uploaded article.
                    
                    **ARTICLE TEXT:**
                    {full_text}
                    
                    {image_context}
                    
                    **TARGET SLIDES:** {slide_count}
                    
                    **REQUIRED STRUCTURE (Journal Club Format):**
                    1.  **Title Slide**: Title, Authors, Journal, Year.
                    2.  **Introduction**: Background, Problem Statement, Hypothesis/Aim.
                    3.  **Methods**: Study Design. *Use layout='procedure_slide' if showing steps or multiple images.*
                    4.  **Results**: Key Findings. *Use layout='results_slide' to MAXIMIZE Tables/Figures for visibility.*
                    5.  **Discussion**: Interpretation of results.
                    6.  **CRITICAL APPRAISAL (Crucial)**:
                        -   Strengths of the study.
                        -   Limitations/Weaknesses (Bias, Confounders, etc.).
                        -   Validity of conclusions.
                    7.  **Conclusion**: Clinical Relevance/Take-home message.
                    8.  **REFERENCES**: List all references used/cited in the presentation. (FINAL SLIDE)
                    
                    **OUTPUT FORMAT:**
                    Return a VALID JSON object with this structure:
                    {{
                      "title": "Presentation Title",
                      "slides": [
                        {{
                          "title": "Slide Title",
                          "layout": "standard", // OPTIONS: "standard", "results_slide" (Maximized Image), "procedure_slide" (2 Images side-by-side)
                          "content": ["Bullet 1", "Bullet 2", "Bullet 3"],
                          "speaker_notes": "Notes for the presenter...",
                          "suggested_image_type": "research_image", // Use 'research_image' generic, OR:
                          "suggested_image_index": 0 // INTEGER. Refer to specific image index from list above.
                        }}
                      ]
                    }}
                    """
                    
                    content_json = call_ai_api(prompt, provider, api_key, model_name=selected_model, json_mode=True)
                    
                    # 3. Generate PPT
                    # We reuse the existing ppt_generator logic
                    ppt_io = generate_presentation(content_json, images_found, temp_dir)
                    
                    st.success(f"Journal Club Presentation Generated! ({len(images_found)} images found)")
                    st.download_button(
                        label="Download Presentation (.pptx)",
                        data=ppt_io,
                        file_name="Journal_Club_Presentation.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation"
                    )
                    
            except Exception as e:
                st.error(f"Error: {e}")

    elif not api_key:
        st.warning("Please enter your API Key.")

    elif not api_key:
         st.warning("Please enter your API Key to proceed.")
